#!/usr/bin/env python3
"""生成 商务 AI Agent 系统开发需求规格说明书 V0.2（最终完整版）"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ═══════════════════════════════════════
# 标题页
# ═══════════════════════════════════════
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER; t.paragraph_format.space_before = Pt(120)
r = t.add_run('商务 AI Agent 系统开发需求规格说明书'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run('第二版 / V0.2（最终版）'); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)
info = doc.add_paragraph(); info.alignment = WD_ALIGN_PARAGRAPH.CENTER; info.paragraph_format.space_before = Pt(40)
r = info.add_run('基于 2026-06-09 会议 + 2026-06-27 商务部深度调研\n2026-06-27'); r.font.size = Pt(11)
doc.add_page_break()

# ═══════════════════════════════════════
# 说明
# ═══════════════════════════════════════
doc.add_heading('说明', level=1)
doc.add_paragraph('本文档为第二版需求规格说明书（V0.2 最终版），基于以下输入整理：')
for item in ['2026-06-09 项目启动会议逐字稿', '2026-06-27 与商务部员工 B 的深度业务调研对话（1272行）', '海外电商/独立站 CRM 销售订单物料-物流 Excel 附件格式', 'CFO 关于跨主体结算和库存归属的讲话精神']:
    doc.add_paragraph(item, style='List Bullet')

items = [
    'ERP 自动制单（P0：一期核心交付）: Save → Submit → Audit',
    '订单类型自动识别（销售/备货）, 不同预审规则链',
    '库存预审三步判断：主体仓库→其他仓库有货→全缺通知销售重新提交',
    'Excel 库存导入页面 + 主体-仓库映射配置管理',
    '产品物料与价格维护页面（备货订单自动计价）',
    '海外电商 Excel 附件兼容（物料-物流每行一个收货人）',
    '中台订单号连续不跳号规则 + 发货通知国内带号/海外不带号',
    '跨主体调货记录与通知（一期标记+通知，二期自动执行）',
    '商务审核前置条件（CRM 合同审批状态检查）',
    '物料多别名匹配（LLM 语义匹配 + 管理台别名维护）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# 修订记录
doc.add_heading('文档修订记录', level=1)
table = doc.add_table(rows=3, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['版本', '日期', '作者', '修订说明']):
    table.rows[0].cells[i].text = h
rs = [('V0.1', '2026-06-09', 'AI整理', '根据会议逐字稿整理第一版'),
      ('V0.2', '2026-06-27', 'AI整理', '补充10项新增需求，修正优先级，补充库存预审/主体-仓库/海外Excel附件等完整设计')]
for ri, (a, b, c, d) in enumerate(rs, 1):
    table.rows[ri].cells[0].text = a; table.rows[ri].cells[1].text = b
    table.rows[ri].cells[2].text = c; table.rows[ri].cells[3].text = d

doc.add_page_break()

# ═══════════════════════════════════════
# 1. 项目背景
# ═══════════════════════════════════════
doc.add_heading('1. 项目背景与会议共识', level=1)
doc.add_paragraph('本项目拟建设"商务 AI Agent 订单中台"，作为 CRM、OMS、金蝶 ERP 之间的流程汇集点和自动化审核中心，替代现有商务人员手工操作流程。')

doc.add_heading('1.1 核心共识', level=2)
for c in [
    '取消人工邮件作为主流程入口，销售订单以 CRM 审批完成数据为源头',
    '中台是订单编排核心：CRM(源头) → 中台(预审) → ERP(制单) → OMS(发货执行)',
    '一期优先解决睿数/渠道订单，逐步扩展到备货和海外订单',
    '统一主数据：物料别名、客户映射、主体-仓库映射必须维护',
    'AI 自动制单(Save→Submit→Audit)是一期核心交付',
]:
    doc.add_paragraph(c, style='List Bullet')

# 当前核心问题
doc.add_heading('1.2 当前核心问题', level=2)
table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['问题类别', '现状', '影响']):
    table.rows[0].cells[i].text = h
for ri, (a, b, c) in enumerate([
    ('重复录入', '同一订单在 CRM、ERP、OMS 多次手工录入', '效率低、错误率高'),
    ('物料名称不对应', 'CRM 只有物料名称无料号，无法区分版本和规格', '商务需人工确认具体料号'),
    ('海外英文订单', '海外订单经常发英文描述，商务需百度翻译', '"语数英"痛点，耗时巨大'),
    ('库存判断靠人工', '商务需多渠道确认库存、沟通调货方案', '每个订单多轮沟通'),
    ('跨主体无记录', '深圳→香港调货无人走流程，财务事后才发现', '对账困难，信息滞后'),
    ('邮件群发全员', '发货通知邮件公布给所有人', '信息过载，相关人漏看'),
    ('范围边界不清', '渠道、电商、备货、C端订单混在一起', '需求扩散'),
], 1):
    table.rows[ri].cells[0].text = a; table.rows[ri].cells[1].text = b; table.rows[ri].cells[2].text = c

doc.add_page_break()

# ═══════════════════════════════════════
# 2. 项目目标与范围
# ═══════════════════════════════════════
doc.add_heading('2. 项目目标、范围与阶段规划', level=1)
doc.add_heading('2.1 总体目标', level=2)
doc.add_paragraph('建立以 CRM 审批完成订单为源头、AI Agent 为中台核心、ERP自动制单+OMS 发货执行的统一业务闭环。')

doc.add_heading('2.2 一期建设目标', level=2)
for g in [
    'CRM 审批完成订单自动进入中台，订单类型自动识别（销售/备货）',
    '中台完成字段完整性、物料别名匹配、客户映射、合同审批状态、库存三步判断的自动预审',
    'ERP 自动制单（Save → Submit → Audit），国内仓发货通知携带中台连续单号',
    '海外电商 Excel 附件自解析：物料-物流合并，每行一个收货人',
    '库存 Excel 导入页面 + 主体-仓库映射配置 + 产品价格维护页面',
    '跨主体调货记录（一期标记+通知，二期自动执行）',
    'OMS 发货状态回流 + 订单追踪归档',
]:
    doc.add_paragraph(g, style='List Bullet')

# 分期范围
doc.add_heading('2.3 分期范围规划', level=2)
table = doc.add_table(rows=3, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['阶段', '目标范围', '主要能力', '不包含/延后']):
    table.rows[0].cells[i].text = h
for ri, (a, b, c, d) in enumerate([
    ('一期内测', '睿数/渠道销售+备货订单\n深圳主体为主', 'CRM同步、预审规则、库存三步判断、ERP自动制单\nExcel附件解析、发货通知、OMS下推\n库存导入、主体-仓库映射、产品价格维护', '跨主体自动拆单、自动结算、自动下推'),
    ('一期扩展', '海外电商/独立站\n多主体覆盖', '主体-仓库映射启用、跨主体调货记录\n海外Excel附件兼容', '跨主体自动结算、自动下推'),
    ('二期', '完整多主体\n自动结算+下推', '跨主体自动拆单、内部调拨单\n自动下推、物流轨迹、财务核验', ''),
], 1):
    for ci, v in enumerate([a, b, c, d]):
        table.rows[ri].cells[ci].text = v

doc.add_page_break()

# ═══════════════════════════════════════
# 3. 核心业务流程
# ═══════════════════════════════════════
doc.add_heading('3. 核心业务流程', level=1)
doc.add_paragraph('完整流程见 docs/workflow-flowchart.png')
doc.add_paragraph('')
p = doc.add_paragraph()
r = p.add_run('流程图（Mermaid）：')
r.bold = True
doc.add_paragraph('docs/workflow-flowchart.md（源码）', style='List Bullet')
doc.add_paragraph('docs/mermaid-chart.html（交互视图）', style='List Bullet')

doc.add_paragraph('')

doc.add_heading('3.1 全流程概览（文字版）', level=2)
steps = [
    'CRM 审批完成订单 → 中台同步（幂等入库，一期范围过滤）',
    '订单类型自动识别（销售/备货）: 金额>0+有附件=销售，金额=0+无附件=备货',
    '前置条件检查（仅销售订单）: CRM合同审批状态=已完成？销售邮箱≠空？',
    '中台建单 MP-{年份}{序号}: 序号连续不跳号，年度重置',
    '11条预审规则链（销售走全量，备货跳过合同/金额/附件规则）',
    '库存三步判断: 主体仓库→其他主体仓库有货→全缺通知销售重新提交',
    'ERP 自动制单: Save → Submit → Audit（金蝶生成 FBillNo）',
    '发货通知邮件: 国内仓发→带中台单号, 海外仓发→不带单号',
    'OMS 下推: 吉客云 API wms.order.create 指数退避重试',
    'OMS 状态追踪 → 订单归档',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'Step {i}：{s}', style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════
# 4. 功能需求规格
# ═══════════════════════════════════════
doc.add_heading('4. 功能需求规格', level=1)
doc.add_paragraph('优先级说明：P0 为一期必须完成，P1 为一期建议/二期优先，P2 为后续增强。')

reqs = [
    ('FR-001', 'CRM 订单同步', '通过 CRM API/webhook 获取审批完成订单，含订单头/明细/客户/附件', 'P0'),
    ('FR-002', '幂等与去重', 'payload_hash 判断，同一订单不重复创建', 'P0'),
    ('FR-003', '订单类型自动识别', '根据金额+附件自动区分销售/备货订单，走不同预审规则链', 'P0'),
    ('FR-004', '前置条件检查', 'CRM 合同审批状态=已完成？销售邮箱？一期范围？阻断时通知销售', 'P0'),
    ('FR-005', '中台订单号生成', 'MP-{年份}{序号} 格式，全局/年度连续不跳号', 'P0'),
    ('FR-006', '物料别名匹配', 'CRM名称→金蝶料号，规则匹配+LLM语义匹配，低置信度进异常', 'P0'),
    ('FR-007', '库存三步判断', 'Step1:主体仓库→Step2:其他仓库→Step3:全缺通知销售重新提交', 'P0'),
    ('FR-008', '主体-仓库映射管理', '管理台配置：主体→关联仓库列表，支持按主体查询库存', 'P0'),
    ('FR-009', '库存Excel导入', '商务下载模板→上传→预览→确认，支持多仓库', 'P0'),
    ('FR-010', '产品价格维护', '每个SKU内部成本价，备货订单自动计价，支持批量导入', 'P0'),
    ('FR-011', 'ERP 自动制单', '预审通过后自动 Save→Submit→Audit，失败进异常', 'P0'),
    ('FR-012', '发货通知', '支持一单多收货人(Excel每行一个)，国内仓发带中台单号', 'P0'),
    ('FR-013', 'OMS 下推', '吉客云 wms.order.create，指数退避重试 60s/180s/540s', 'P0'),
    ('FR-014', 'OMS 状态回写', '拣货/发货/签收状态轮询更新', 'P0'),
    ('FR-015', '订单变更跟踪', 'payload_hash 检测CRM变更，区分已闭环/未闭环/部分发货3种处理', 'P0'),
    ('FR-016', '跨主体调货记录', '订单主体≠出货主体时生成调货记录，通知双方财务', 'P1'),
    ('FR-017', '特殊需求分类', 'LLM提取，按物流/生产/报关分类，分流到对应部门', 'P1'),
    ('FR-018', '备货订单完整支持', '跳过合同/金额/附件规则，价格从产品表取', 'P0'),
    ('FR-019', '海外英文订单', 'LLM 自动翻译+语义匹配物料', 'P1'),
    ('FR-020', '异常任务管理', '创建/分派/处理/关闭/重开/SLA', 'P0'),
    ('FR-021', '自动通知', '按角色按需通知（不是全员群发）', 'P0'),
    ('FR-022', '审计日志', '所有操作留痕可追溯', 'P0'),
    ('FR-023', '跨主体自动结算', '生成内部调拨单，推送给双方财务（二期）', 'P2'),
    ('FR-024', '自动下推', '物流发货后自动标记金蝶出库（二期）', 'P2'),
]

table = doc.add_table(rows=len(reqs)+1, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['ID', '功能模块', '需求描述', '优先级']):
    table.rows[0].cells[i].text = h
for ri, (id_, mod, desc, pri) in enumerate(reqs, 1):
    table.rows[ri].cells[0].text = id_
    table.rows[ri].cells[1].text = mod
    table.rows[ri].cells[2].text = desc
    table.rows[ri].cells[3].text = pri

doc.add_page_break()

# ═══════════════════════════════════════
# 5. 预审规则细则
# ═══════════════════════════════════════
doc.add_heading('5. 预审规则细则', level=1)

table = doc.add_table(rows=12, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['规则ID', '规则名称', '销售订单', '备货订单']):
    table.rows[0].cells[i].text = h
for ri, (id_, name, sale, stock) in enumerate([
    ('BR-001', '必填字段校验', '客户/物料/数量/金额/收货信息', '客户/物料/数量'),
    ('BR-002', '一期范围检查', '检查', '检查'),
    ('BR-003', '客户主数据映射', '检查', '检查'),
    ('BR-004', '金额正数', '检查', '跳过'),
    ('BR-005', '金额一致性', '检查(CRM vs 附件)', '跳过'),
    ('BR-006', '物料匹配(别名)', '规则+LLM语义匹配', '规则+LLM语义匹配'),
    ('BR-007', '合同金额一致性', '检查', '跳过'),
    ('BR-008', '附件一致性', '检查(PI vs 订单)', '跳过'),
    ('BR-009', '库存三步判断', '主体→其他→全缺', '主体→其他→全缺'),
    ('BR-010', '价格校验', 'CRM价格>0', '取自产品价格表'),
    ('BR-011', '商务审核前置', 'CRM合同审批状态', '跳过'),
], 1):
    table.rows[ri].cells[0].text = id_
    table.rows[ri].cells[1].text = name
    table.rows[ri].cells[2].text = sale
    table.rows[ri].cells[3].text = stock

doc.add_page_break()

# ═══════════════════════════════════════
# 6. 库存预审详细设计
# ═══════════════════════════════════════
doc.add_heading('6. 库存预审详细设计', level=1)

doc.add_heading('6.1 数据来源', level=2)
doc.add_paragraph('库存数据来源：商务每周通过涛哥（仓管）获取最新库存表，通过中台Excel导入功能更新。', style='List Bullet')
doc.add_paragraph('库存表字段：仓库、物料编码、物料名称、库存数量。不分主体（仓库不区分主体）。', style='List Bullet')

doc.add_heading('6.2 主体-仓库映射', level=2)
doc.add_paragraph('管理台配置：主体编码→关联仓库列表。如：', style='List Bullet')
doc.add_paragraph('HK(香港) → 欧洲仓、美国仓', style='List Bullet')
doc.add_paragraph('SZ(深圳) → 武汉仓', style='List Bullet')
doc.add_paragraph('LU(卢森堡) → 欧洲仓', style='List Bullet')
doc.add_paragraph('US(美国) → 美国仓', style='List Bullet')
doc.add_paragraph('注：一个仓库可归属多个主体（如欧洲仓同时归属HK和LU）', style='List Bullet')

doc.add_heading('6.3 三步判断逻辑', level=2)
steps_desc = [
    ('Step 1：主体对应仓库检查', '根据PI/CRM解析的销售主体 → 查主体关联仓库的库存 → 够→正常走，不够→Step2'),
    ('Step 2：其他主体仓库检查', '查其他主体关联仓库是否有库存 → 有→标记"可调货"，通知B确认 → 无→Step3'),
    ('Step 3：全缺货通知销售', '通知CRM订单的销售负责人，告知缺货物料和可能有货的仓库信息 → 销售协调后重新在CRM提交'),
]
for title, desc in steps_desc:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}：'); r.bold = True
    p.add_run(desc)

doc.add_page_break()

# ═══════════════════════════════════════
# 7. 外部系统接口
# ═══════════════════════════════════════
doc.add_heading('7. 外部系统接口需求', level=1)

table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['系统', '接口方向', '说明']):
    table.rows[0].cells[i].text = h
for ri, (a, b, c) in enumerate([
    ('CRM(纷享销客)', '读', '审批完成订单/明细/客户/附件/PI 同步'),
    ('金蝶云星空(ERP)', '读写', '查询物料/库存/客户；自动 Save→Submit→Audit 销售订单'),
    ('吉客云(OMS)', '读写', '创建发货单(wms.order.create)，状态回写'),
    ('邮件(腾讯企业邮)', '发送通知', '发货通知/异常通知/调货通知，按需发送'),
], 1):
    table.rows[ri].cells[0].text = a; table.rows[ri].cells[1].text = b; table.rows[ri].cells[2].text = c

doc.add_page_break()

# ═══════════════════════════════════════
# 8. AI Agent 要求
# ═══════════════════════════════════════
doc.add_heading('8. AI Agent 设计要求', level=1)
doc.add_paragraph('AI Agent 定位：嵌入中台的"流程编排+审单解释+异常协同"智能助手。', style='List Bullet')
doc.add_paragraph('物料语义匹配：CRM名称→金蝶料号，用于别名库无法命中的场景。', style='List Bullet')
doc.add_paragraph('海外英文订单翻译+语义匹配：解决商务"语数英"痛点。', style='List Bullet')
doc.add_paragraph('特殊需求提取+分类：从附件/备注中提取非标要求，按物流/生产/报关分流。', style='List Bullet')
doc.add_paragraph('异常解释：规则阻断时AI生成自然语言解释+补正建议。', style='List Bullet')
doc.add_paragraph('所有AI输出可追溯（输入摘要+Prompt版本+模型输出+置信度+人工反馈）。', style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════
# 9. 管理台功能
# ═══════════════════════════════════════
doc.add_heading('9. 一期管理台新增功能', level=1)

table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['功能', '说明', '优先级']):
    table.rows[0].cells[i].text = h
for ri, (a, b, c) in enumerate([
    ('库存Excel导入', '下载模板→上传→预览→确认，覆盖该仓库旧数据', 'P0'),
    ('库存查询', '按仓库/物料查看当前快照+最近更新时间', 'P0'),
    ('主体-仓库映射', '主体编码→关联仓库列表，支持增删改', 'P0'),
    ('产品价格维护', '每条SKU的内部成本价，支持批量导入', 'P0'),
    ('产品别名维护', '每条SPU的业务别号(CRM名称映射)', 'P0'),
    ('ERP配置页面', '金蝶连接参数+写入开关(erp_write_enabled)', 'P0'),
], 1):
    table.rows[ri].cells[0].text = a; table.rows[ri].cells[1].text = b; table.rows[ri].cells[2].text = c

doc.add_page_break()

# ═══════════════════════════════════════
# 10. 测试与验收
# ═══════════════════════════════════════
doc.add_heading('10. 测试与验收标准', level=1)
table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['ID', '场景', '标准']):
    table.rows[0].cells[i].text = h
for ri, (a, b, c) in enumerate([
    ('AC-001', 'ERP制单全链路', 'Save→Submit→Audit 一站式通过，测试完自动删除'),
    ('AC-002', '库存三步判断', '主体仓库够→正常；其他有货→通知B；全缺→通知销售'),
    ('AC-003', '订单类型识别', '有金额+附件=销售；无金额=备货；走不同规则链'),
    ('AC-004', 'Excel附件解析', '5行3收货人→预审正确→发货通知含多收货人'),
    ('AC-005', '单号连续性', 'MP-20270001→0002→0003...连续不跳号'),
    ('AC-006', '发货通知邮件', '国内仓发→带单号；海外仓发→不带单号'),
    ('AC-007', '备货订单全流程', 'CRM下单(无价格无附件)→中台识别→预审(跳过合同/金额规则)→产品表取价→ERP制单'),
], 1):
    table.rows[ri].cells[0].text = a; table.rows[ri].cells[1].text = b; table.rows[ri].cells[2].text = c

doc.add_page_break()

# 保存
out = '/sessions/amazing-nifty-gauss/mnt/jm-sp-bot/docs/商务AI_Agent_系统开发需求规格说明书_v0.2.docx'
doc.save(out)
print(f'OK: {out}')
