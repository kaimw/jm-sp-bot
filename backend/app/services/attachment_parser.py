from __future__ import annotations

import io
import json
import os
import re
import subprocess
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from PIL import Image


ATTACHMENT_TEXT_PARSER_VERSION = "paddle-structure-ocr-v2"
_PADDLE_OCR_ENGINE = None
_PADDLE_STRUCTURE_ENGINE = None


@dataclass
class ParsedAttachment:
    file_name: str
    status: str
    text: str = ""
    error: str | None = None
    archive_path: str | None = None
    archive_depth: int | None = None
    children: list["ParsedAttachment"] = field(default_factory=list)
    metadata: dict[str, str] = field(default_factory=dict)


def parse_attachment(file_name: str, content: bytes, *, max_zip_bytes: int, max_depth: int, depth: int = 0) -> ParsedAttachment:
    suffix = Path(file_name).suffix.lower()
    try:
        if suffix == ".docx":
            return ParsedAttachment(file_name=file_name, status="Parsed", text=parse_docx(content), archive_depth=depth)
        if suffix == ".xlsx":
            return ParsedAttachment(file_name=file_name, status="Parsed", text=parse_xlsx(content), archive_depth=depth)
        if suffix == ".zip":
            return parse_zip(file_name, content, max_zip_bytes=max_zip_bytes, max_depth=max_depth, depth=depth)
        if suffix in {".txt", ".csv"}:
            return ParsedAttachment(file_name=file_name, status="Parsed", text=content.decode("utf-8", errors="replace"), archive_depth=depth)
        if suffix == ".pdf":
            text, metadata = parse_pdf_text_with_metadata(content)
            return ParsedAttachment(file_name=file_name, status="Parsed", text=text, archive_depth=depth, metadata=metadata)
        if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
            text, metadata = parse_image_text_with_metadata(content, suffix)
            return ParsedAttachment(file_name=file_name, status="Parsed", text=text, archive_depth=depth, metadata=metadata)
        return ParsedAttachment(file_name=file_name, status="Skipped", error=f"Unsupported file type: {suffix or 'unknown'}", archive_depth=depth)
    except Exception as exc:
        return ParsedAttachment(file_name=file_name, status="Failed", error=str(exc), archive_depth=depth)


def parse_docx(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    parts.extend(extra for extra in parse_docx_document_xml_text(content) if extra.strip())
    return dedupe_text_lines(parts)


def parse_docx_document_xml_text(content: bytes) -> list[str]:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            xml = archive.read("word/document.xml")
    except Exception:
        return []
    try:
        root = ET.fromstring(xml)
    except ET.ParseError:
        return []
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    lines: list[str] = []
    for paragraph in root.iter(f"{ns}p"):
        chunks: list[str] = []
        for node in paragraph.iter():
            if node.tag == f"{ns}t" and node.text:
                chunks.append(node.text)
            elif node.tag == f"{ns}tab":
                chunks.append(" ")
            elif node.tag == f"{ns}br":
                chunks.append("\n")
        text = repair_mojibake_text("".join(chunks))
        text = re.sub(r"[ \t\r\n]+", " ", text).strip()
        if text:
            lines.append(text)
    return lines


def dedupe_text_lines(lines: list[str]) -> str:
    result: list[str] = []
    seen: set[str] = set()
    for line in lines:
        text = str(line or "").strip()
        if not text:
            continue
        key = re.sub(r"\s+", "", text)
        if key in seen:
            continue
        seen.add(key)
        result.append(text)
    return "\n".join(result)


def parse_xlsx(content: bytes) -> str:
    workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"[{sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value not in (None, "")]
            if values:
                parts.append(" | ".join(values))
    workbook.close()
    return "\n".join(parts)


def parse_pdf_text(content: bytes) -> str:
    return parse_pdf_text_with_metadata(content)[0]


def parse_pdf_text_with_metadata(content: bytes) -> tuple[str, dict[str, str]]:
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(io.BytesIO(content))
        parts = []
        for index, page in enumerate(reader.pages, start=1):
            text = repair_mojibake_text(page.extract_text() or "")
            if text.strip():
                parts.append(f"[page {index}]\n{text.strip()}")
        if parts:
            extracted = "\n\n".join(parts)
            if is_meaningful_text(extracted):
                return extracted, {"parser_version": ATTACHMENT_TEXT_PARSER_VERSION, "ocr_engine": "native_pdf_text"}
    except ModuleNotFoundError:
        pass
    except Exception:
        pass

    text = content.decode("latin-1", errors="ignore")
    matches = re.findall(r"\(([^()]*)\)\s*Tj", text)
    if not matches:
        matches = re.findall(r"\(([^()]*)\)", text)
    decoded = [decode_pdf_literal(match).strip() for match in matches]
    result = "\n".join(item for item in decoded if item)
    if result and is_meaningful_text(result):
        return result, {"parser_version": ATTACHMENT_TEXT_PARSER_VERSION, "ocr_engine": "pdf_literal_text"}
    ocr_text, metadata = parse_pdf_by_ocr_with_metadata(content)
    if ocr_text:
        return ocr_text, metadata
    return result, {"parser_version": ATTACHMENT_TEXT_PARSER_VERSION, "ocr_engine": "none"}


def is_meaningful_text(value: str) -> bool:
    text = str(value or "").strip()
    if len(text) < 20:
        return bool(text)
    readable = sum(1 for char in text if char.isalnum() or "\u4e00" <= char <= "\u9fff")
    controls = sum(1 for char in text if ord(char) < 32 and char not in "\n\r\t")
    return readable / max(1, len(text)) >= 0.08 and controls / max(1, len(text)) <= 0.02


def parse_image_text(content: bytes, suffix: str = ".png") -> str:
    return parse_image_text_with_metadata(content, suffix)[0]


def parse_image_text_with_metadata(content: bytes, suffix: str = ".png") -> tuple[str, dict[str, str]]:
    if should_use_ocr_worker():
        return parse_image_text_with_worker(content, suffix)
    return parse_image_text_in_process_with_metadata(content, suffix)


def should_use_ocr_worker() -> bool:
    if os.getenv("ATTACHMENT_OCR_IN_WORKER") == "1":
        return False
    return os.getenv("ATTACHMENT_OCR_USE_WORKER", "true").strip().lower() not in {"0", "false", "no", "off"}


def parse_image_text_with_worker(content: bytes, suffix: str = ".png") -> tuple[str, dict[str, str]]:
    timeout = int(os.getenv("ATTACHMENT_OCR_WORKER_TIMEOUT", "180"))
    with tempfile.NamedTemporaryFile(suffix=suffix) as input_file, tempfile.NamedTemporaryFile(suffix=".json") as output_file:
        input_file.write(content)
        input_file.flush()
        env = os.environ.copy()
        env["ATTACHMENT_OCR_IN_WORKER"] = "1"
        command = [
            sys.executable,
            "-m",
            "backend.app.services.ocr_worker",
            "--input",
            input_file.name,
            "--suffix",
            suffix,
            "--output",
            output_file.name,
        ]
        completed = subprocess.run(command, capture_output=True, text=True, timeout=timeout, check=False, env=env)
        output = read_ocr_worker_output(output_file.name)
    if output.get("ok") and str(output.get("text") or "").strip():
        metadata = output.get("metadata") if isinstance(output.get("metadata"), dict) else {}
        metadata = {str(key): str(value) for key, value in metadata.items()}
        metadata["ocr_worker"] = "subprocess"
        metadata["ocr_worker_exit_code"] = str(completed.returncode)
        return str(output.get("text") or ""), metadata
    error = str(output.get("error") or "").strip()
    stderr = (completed.stderr or "").strip()
    stdout = (completed.stdout or "").strip()
    detail = "; ".join(item for item in (error, stderr[-1000:], stdout[-1000:]) if item)
    raise ValueError(detail or f"OCR worker failed with exit code {completed.returncode}")


def read_ocr_worker_output(path: str) -> dict:
    try:
        with open(path, encoding="utf-8") as handle:
            data = json.load(handle)
    except Exception:
        return {}
    return data if isinstance(data, dict) else {}


def parse_image_text_in_process_with_metadata(content: bytes, suffix: str = ".png") -> tuple[str, dict[str, str]]:
    image = Image.open(io.BytesIO(content))
    if image.mode not in {"RGB", "L"}:
        image = image.convert("RGB")
    errors: list[str] = []
    for engine in configured_ocr_engines():
        if engine == "paddle_structure":
            try:
                text = run_paddle_structure(image, suffix=suffix)
                if text.strip():
                    return text, {"parser_version": ATTACHMENT_TEXT_PARSER_VERSION, "ocr_engine": "paddle_structure"}
            except Exception as exc:
                errors.append(f"paddle_structure:{str(exc)[:200]}")
        if engine == "paddle":
            try:
                text = run_paddleocr(image, suffix=suffix)
                if text.strip():
                    return text, {"parser_version": ATTACHMENT_TEXT_PARSER_VERSION, "ocr_engine": "paddleocr"}
            except Exception as exc:
                errors.append(f"paddleocr:{str(exc)[:200]}")
    if errors:
        raise ValueError("; ".join(errors))
    return "", {"parser_version": ATTACHMENT_TEXT_PARSER_VERSION, "ocr_engine": "none"}


def parse_pdf_by_ocr(content: bytes, *, max_pages: int = 3) -> str:
    return parse_pdf_by_ocr_with_metadata(content, max_pages=max_pages)[0]


def parse_pdf_by_ocr_with_metadata(content: bytes, *, max_pages: int = 3) -> tuple[str, dict[str, str]]:
    try:
        import fitz  # type: ignore
    except ModuleNotFoundError:
        return "", {"parser_version": ATTACHMENT_TEXT_PARSER_VERSION, "ocr_engine": "none"}
    parts: list[str] = []
    engines: list[str] = []
    document = fitz.open(stream=content, filetype="pdf")
    try:
        for page_index in range(min(max_pages, len(document))):
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
            text, metadata = parse_image_text_with_metadata(pixmap.tobytes("png"), ".png")
            engines.append(metadata.get("ocr_engine", "unknown"))
            if text.strip():
                parts.append(f"[ocr page {page_index + 1}]\n{text.strip()}")
    finally:
        document.close()
    return "\n\n".join(parts), {
        "parser_version": ATTACHMENT_TEXT_PARSER_VERSION,
        "ocr_engine": ",".join(sorted(set(engine for engine in engines if engine))) or "none",
    }


def configured_ocr_engines() -> list[str]:
    raw = os.getenv("ATTACHMENT_OCR_ENGINES", "paddle_structure,paddle")
    engines = [item.strip().lower() for item in re.split(r"[,;]+", raw) if item.strip()]
    normalized = ["paddle" if engine == "paddleocr" else engine for engine in engines]
    return [engine for engine in normalized if engine in {"paddle_structure", "paddle"}]


def run_paddle_structure(image: Image.Image, *, suffix: str = ".png") -> str:
    global _PADDLE_STRUCTURE_ENGINE
    with tempfile.NamedTemporaryFile(suffix=suffix) as handle:
        image.save(handle.name)
        if _PADDLE_STRUCTURE_ENGINE is None:
            try:
                from paddleocr import PPStructure  # type: ignore

                try:
                    _PADDLE_STRUCTURE_ENGINE = PPStructure(show_log=False, lang="ch")
                except TypeError:
                    _PADDLE_STRUCTURE_ENGINE = PPStructure(lang="ch")
            except ImportError:
                from paddleocr import PPStructureV3  # type: ignore

                _PADDLE_STRUCTURE_ENGINE = PPStructureV3(
                    lang="ch",
                    use_table_recognition=True,
                    use_formula_recognition=False,
                    use_chart_recognition=False,
                    use_region_detection=False,
                    use_seal_recognition=False,
                )
        engine = _PADDLE_STRUCTURE_ENGINE
        if hasattr(engine, "predict"):
            result = engine.predict(handle.name)
        else:
            result = engine(handle.name)
    lines = paddle_structure_lines(result)
    return "\n".join(line for line in lines if line.strip())


def paddle_structure_lines(result) -> list[str]:
    lines: list[str] = []

    def visit(value) -> None:
        if value is None:
            return
        if isinstance(value, dict):
            item_type = str(value.get("type") or value.get("label") or "").lower()
            res = value.get("res") if "res" in value else value.get("result")
            if item_type == "table":
                html = ""
                if isinstance(res, dict):
                    html = str(res.get("html") or res.get("html_content") or "")
                elif isinstance(res, str):
                    html = res
                rows = html_table_to_pipe_rows(html)
                if rows:
                    lines.extend(rows)
                    return
            if isinstance(res, list):
                for item in res:
                    visit(item)
                return
            text = value.get("text") or value.get("content")
            if isinstance(text, str) and text.strip():
                lines.append(text.strip())
                return
            for nested in value.values():
                visit(nested)
            return
        if isinstance(value, list):
            for item in value:
                visit(item)
            return
        if isinstance(value, str) and value.strip():
            text = value.strip()
            if "<table" in text.lower():
                rows = html_table_to_pipe_rows(text)
                if rows:
                    lines.extend(rows)
                    return
            lines.append(text)

    visit(result)
    return dedupe_text_lines(lines).splitlines()


class _SimpleTableParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.rows: list[list[str]] = []
        self._row: list[str] | None = None
        self._cell: list[str] | None = None

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag == "tr":
            self._row = []
        elif tag in {"td", "th"} and self._row is not None:
            self._cell = []

    def handle_data(self, data: str) -> None:
        if self._cell is not None:
            self._cell.append(data)

    def handle_endtag(self, tag: str) -> None:
        if tag in {"td", "th"} and self._row is not None and self._cell is not None:
            self._row.append(re.sub(r"\s+", " ", "".join(self._cell)).strip())
            self._cell = None
        elif tag == "tr" and self._row is not None:
            if any(cell for cell in self._row):
                self.rows.append(self._row)
            self._row = None


def html_table_to_pipe_rows(html: str) -> list[str]:
    if not str(html or "").strip():
        return []
    parser = _SimpleTableParser()
    parser.feed(str(html))
    return [" | ".join(cell for cell in row if cell) for row in parser.rows if row]


def run_paddleocr(image: Image.Image, *, suffix: str = ".png") -> str:
    global _PADDLE_OCR_ENGINE
    with tempfile.NamedTemporaryFile(suffix=suffix) as handle:
        image.save(handle.name)
        if _PADDLE_OCR_ENGINE is None:
            from paddleocr import PaddleOCR  # type: ignore

            try:
                _PADDLE_OCR_ENGINE = PaddleOCR(lang="ch", use_textline_orientation=True)
            except TypeError:
                try:
                    _PADDLE_OCR_ENGINE = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
                except TypeError:
                    _PADDLE_OCR_ENGINE = PaddleOCR(lang="ch")
        engine = _PADDLE_OCR_ENGINE
        if hasattr(engine, "predict"):
            result = engine.predict(handle.name)
        else:
            try:
                result = engine.ocr(handle.name, cls=True)
            except TypeError:
                result = engine.ocr(handle.name)
    lines = paddleocr_result_lines(result)
    return "\n".join(line for line in lines if line.strip())


def paddleocr_result_lines(result) -> list[str]:
    lines: list[str] = []

    def visit(value) -> None:
        if value is None:
            return
        if isinstance(value, dict):
            texts = value.get("rec_texts") or value.get("texts") or value.get("text")
            if isinstance(texts, list):
                for item in texts:
                    if str(item or "").strip():
                        lines.append(str(item).strip())
                return
            if isinstance(texts, str) and texts.strip():
                lines.append(texts.strip())
                return
            for nested in value.values():
                visit(nested)
            return
        if isinstance(value, tuple) and len(value) >= 2 and isinstance(value[1], tuple):
            text = value[1][0]
            if str(text or "").strip():
                lines.append(str(text).strip())
            return
        if isinstance(value, list):
            for item in value:
                visit(item)

    visit(result)
    return dedupe_text_lines(lines).splitlines()


def decode_pdf_literal(value: str) -> str:
    decoded = (
        value.replace(r"\(", "(")
        .replace(r"\)", ")")
        .replace(r"\\", "\\")
        .replace(r"\n", "\n")
        .replace(r"\r", "\n")
        .replace(r"\t", "\t")
    )
    try:
        return decoded.encode("latin-1").decode("utf-8")
    except UnicodeError:
        return repair_mojibake_text(decoded)


def repair_mojibake_text(value: str) -> str:
    text = str(value or "")
    if not text:
        return ""
    if any(marker in text for marker in ("Ã", "Â", "å", "æ", "è", "é", "ï¼")):
        for source_encoding in ("latin-1", "cp1252"):
            try:
                raw = text.encode(source_encoding)
            except UnicodeEncodeError:
                continue
            for target_encoding in ("utf-8", "gb18030", "gbk", "big5"):
                try:
                    repaired = raw.decode(target_encoding)
                except UnicodeDecodeError:
                    continue
                if repaired != text and "\ufffd" not in repaired:
                    return repaired
    return text


def parse_zip(file_name: str, content: bytes, *, max_zip_bytes: int, max_depth: int, depth: int) -> ParsedAttachment:
    if len(content) > max_zip_bytes:
        return ParsedAttachment(file_name=file_name, status="Failed", error="ZIP exceeds maximum size.", archive_depth=depth)
    if depth >= max_depth:
        return ParsedAttachment(file_name=file_name, status="Failed", error="ZIP exceeds maximum extraction depth.", archive_depth=depth)

    root = ParsedAttachment(file_name=file_name, status="Parsed", archive_depth=depth)
    total_uncompressed = 0
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        for info in archive.infolist():
            if info.is_dir():
                continue
            total_uncompressed += info.file_size
            if total_uncompressed > max_zip_bytes:
                root.children.append(
                    ParsedAttachment(file_name=info.filename, status="Failed", error="ZIP uncompressed size exceeds limit.", archive_depth=depth + 1)
                )
                continue
            child_content = archive.read(info)
            child = parse_attachment(info.filename, child_content, max_zip_bytes=max_zip_bytes, max_depth=max_depth, depth=depth + 1)
            child.archive_path = info.filename
            root.children.append(child)
    root.text = "\n\n".join(child.text for child in root.children if child.text)
    return root
